import docx
from docx.shared import RGBColor
from datetime import date

def doc_builder(summary_list):

    # Create an instance of a word document
    doc = docx.Document()

    # Add a Title to the document
    doc.add_heading('MEDIA COVERAGE ON SECURITY ISSUES ', 0)

    doc.add_heading('Date: {}'.format(date.today().strftime("%d/%m/%Y")), 1).alignment = 2


    for i, summary in enumerate(summary_list):

        # Add paragraph
        doc.add_heading('Summary {}:'.format(i+1), 3)
        para = doc.add_paragraph().add_run(summary)


    # Save the document to a location
    doc.save('downloads/summary report.docx')
